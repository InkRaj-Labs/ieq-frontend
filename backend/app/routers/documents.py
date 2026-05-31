"""
IEQ Documents & RAG Router
GET    /ieq/documents
POST   /ieq/documents/upload
GET    /ieq/documents/{id}
DELETE /ieq/documents/{id}
GET    /ieq/rag/search
POST   /ieq/rag/chat
"""

import asyncio
import logging
import os
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc

from app.core.config import settings
from app.core.database import get_db, Document
from app.core.events import event_bus
from app.services.embeddings import embed_texts, chunk_text
from app.services.qdrant_service import qdrant_service

router = APIRouter()
logger = logging.getLogger("ieq.documents")

ALLOWED_TYPES = {"application/pdf", "text/plain", "text/markdown",
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


@router.get("/documents")
async def list_documents(db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(Document).order_by(desc(Document.created_at)).limit(100)
    )
    docs = result.scalars().all()
    return {
        "documents": [
            {
                "id": d.id,
                "filename": d.original_name,
                "file_type": d.file_type,
                "file_size": d.file_size,
                "chunk_count": d.chunk_count,
                "status": d.status,
                "created_at": d.created_at.isoformat() if d.created_at else None,
            }
            for d in docs
        ]
    }


@router.post("/documents/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """Upload and process a document for RAG."""
    # Validate
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Size check
    content = await file.read()
    max_bytes = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(status_code=413, detail=f"File too large (max {settings.MAX_UPLOAD_SIZE_MB}MB)")

    # Save to disk
    upload_dir = Path(settings.UPLOAD_DIR) / "documents"
    upload_dir.mkdir(parents=True, exist_ok=True)

    safe_name = f"{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}_{Path(file.filename).name}"
    file_path = upload_dir / safe_name
    file_path.write_bytes(content)

    # DB record
    doc = Document(
        filename=safe_name,
        original_name=file.filename or safe_name,
        file_type=ext.lstrip("."),
        file_size=len(content),
        status="processing",
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)
    doc_id = doc.id

    background_tasks.add_task(_process_document, doc_id, str(file_path), ext)

    await event_bus.emit(
        event_type="upload",
        service_id="qdrant",
        title=f"Uploaded: {file.filename}",
        db=db,
    )

    return {
        "id": doc_id,
        "status": "processing",
        "filename": file.filename,
    }


@router.get("/documents/{document_id}")
async def get_document(document_id: int, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return {
        "id": doc.id,
        "filename": doc.original_name,
        "file_type": doc.file_type,
        "file_size": doc.file_size,
        "chunk_count": doc.chunk_count,
        "status": doc.status,
        "error": doc.error,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
    }


@router.delete("/documents/{document_id}")
async def delete_document(document_id: int, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove from Qdrant
    await qdrant_service.delete_document(document_id)

    # Remove file
    try:
        upload_dir = Path(settings.UPLOAD_DIR) / "documents"
        file_path = upload_dir / doc.filename
        if file_path.exists():
            file_path.unlink()
    except Exception as e:
        logger.warning(f"Could not delete file: {e}")

    await db.delete(doc)
    return {"status": "deleted", "id": document_id}


# ── RAG ───────────────────────────────────────────────────────────────────────

@router.get("/rag/search")
async def rag_search(query: str, limit: int = 5):
    """Semantic search over documents."""
    if not query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    embeddings = await embed_texts([query])
    results = await qdrant_service.search(embeddings[0], limit=limit)

    return {
        "query": query,
        "results": results,
    }


class RAGChatRequest(BaseModel):
    query: str
    conversation_id: Optional[int] = None
    document_ids: Optional[List[int]] = None
    model_key: Optional[str] = None


@router.post("/rag/chat")
async def rag_chat(req: RAGChatRequest, db: AsyncSession = Depends(get_db)):
    """RAG-augmented chat: search docs then answer with context."""
    from app.services.lm_studio import lm_studio

    # Search for relevant chunks
    embeddings = await embed_texts([req.query])
    chunks = await qdrant_service.search(
        embeddings[0],
        limit=5,
        document_ids=req.document_ids,
    )

    context = "\n\n".join(c["text"] for c in chunks) if chunks else "No relevant documents found."

    messages = [
        {
            "role": "system",
            "content": (
                "You are IEQ Assistant. Answer questions based on the provided context.\n\n"
                f"CONTEXT:\n{context}"
            ),
        },
        {"role": "user", "content": req.query},
    ]

    model = req.model_key or settings.DEFAULT_MODEL
    full_response = ""
    async for token in lm_studio.chat_complete(
        messages=messages,
        model=model,
        stream=False,
    ):
        full_response += token

    await event_bus.emit(
        event_type="chat",
        service_id="qdrant",
        title=f"RAG query: {req.query[:50]}",
        db=db,
    )

    return {
        "answer": full_response,
        "sources": chunks,
        "model": model,
    }


# ── Background Processing ─────────────────────────────────────────────────────

async def _process_document(doc_id: int, file_path: str, ext: str):
    """Background: extract text, chunk, embed, store in Qdrant."""
    from app.core.database import AsyncSessionLocal

    async with AsyncSessionLocal() as db:
        result = await db.execute(select(Document).where(Document.id == doc_id))
        doc = result.scalar_one_or_none()
        if not doc:
            return

        try:
            text = await _extract_text(file_path, ext)
            if not text.strip():
                raise ValueError("No text extracted from document")

            chunks = chunk_text(text, chunk_size=500, overlap=50)
            if not chunks:
                raise ValueError("No chunks generated")

            embeddings = await embed_texts(chunks)
            success = await qdrant_service.upsert_chunks(
                document_id=doc_id,
                chunks=chunks,
                embeddings=embeddings,
                metadata={"filename": doc.original_name, "file_type": doc.file_type},
            )

            doc.status = "ready" if success else "failed"
            doc.chunk_count = len(chunks)
            if not success:
                doc.error = "Failed to store in Qdrant"

        except Exception as e:
            logger.error(f"Document {doc_id} processing failed: {e}")
            doc.status = "failed"
            doc.error = str(e)

        await db.commit()


async def _extract_text(file_path: str, ext: str) -> str:
    """Extract text from various file types."""
    loop = asyncio.get_event_loop()

    if ext in (".txt", ".md"):
        return Path(file_path).read_text(encoding="utf-8", errors="ignore")

    elif ext == ".pdf":
        def _read_pdf():
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        return "\n".join(page.extract_text() or "" for page in pdf.pages)
                except ImportError:
                    raise RuntimeError("No PDF library available. Install pypdf or pdfplumber.")
        return await loop.run_in_executor(None, _read_pdf)

    elif ext == ".docx":
        def _read_docx():
            try:
                import docx
                document = docx.Document(file_path)
                return "\n".join(p.text for p in document.paragraphs)
            except ImportError:
                raise RuntimeError("python-docx not installed.")
        return await loop.run_in_executor(None, _read_docx)

    else:
        raise ValueError(f"Unsupported file type: {ext}")
